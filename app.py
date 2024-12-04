import tempfile
import zipfile
import streamlit as st
from src.video_preview import VideoPreview
from src.transcriber import TranscriptionManager
from src.playlist_processor import PlaylistProcessor
from src.styles import CustomCSS
from src.utils import (
    ModelConfig,
    format_size, 
    is_playlist_url, 
    extract_playlist_info, 
    format_duration,
    check_system_compatibility,
    VideoUtility,
    AudioUtility
)
from pathlib import Path
import os
import torch
from datetime import datetime

# Must be first Streamlit command
st.set_page_config(
    page_title="YouTube & Audio Transcriber",
    page_icon="📝",
    layout="wide",
    menu_items={
        'About': "YouTube & Audio Transcriber - Convert videos and audio to text"
    }
)

# Apply custom CSS
st.markdown(CustomCSS.STYLES, unsafe_allow_html=True)

def init_session_state():
    """Initialize session state with default values"""
    if 'initialized' not in st.session_state:
        st.session_state.initialized = True
        st.session_state.manager = TranscriptionManager()
        st.session_state.preview = VideoPreview()
        st.session_state.initialization_checked = False
        
        st.session_state.model_option = "Large"
        st.session_state.model_info = ModelConfig.MODELS["Large"].copy()
        st.session_state.model_key = "large"

def check_initialization():
    """Check system compatibility and requirements"""
    if not st.session_state.initialization_checked:
        compatible, issues = check_system_compatibility()
        
        if not compatible:
            st.error("### ⚠️ System Compatibility Issues")
            for issue in issues:
                st.error(f"- {issue}")
            st.stop()
        elif issues:
            with st.expander("⚠️ System Warnings", expanded=False):
                for issue in issues:
                    st.warning(f"- {issue}")
        
        st.session_state.initialization_checked = True

def render_sidebar():
    """Render sidebar with model settings and system info"""
    with st.sidebar:
        st.header("Model Settings")

        # Model selection
        current_model = st.session_state.model_option
        available_models = list(ModelConfig.MODELS.keys())
        current_index = available_models.index(current_model) if current_model in available_models else 0

        selected_model = st.selectbox(
            "Select Model",
            options=available_models,
            index=current_index,
            key="model_select"
        )

        if selected_model != st.session_state.model_option:
            st.session_state.model_option = selected_model
            st.session_state.model_info = ModelConfig.MODELS[selected_model].copy()
            st.session_state.model_key = ModelConfig.MODELS[selected_model]["key"]
            st.rerun()

        # Model information display
        st.markdown("### Model Details")
        info_cols = st.columns(2)
        with info_cols[0]:
            st.metric("Size", st.session_state.model_info['size'])
            st.metric("Speed", f"{st.session_state.model_info['speed_factor']}x")
        with info_cols[1]:
            st.metric("VRAM", st.session_state.model_info['memory_requirement'])
            st.metric("Download", st.session_state.model_info['download_size'])
        
        st.markdown(f"**Description:** {st.session_state.model_info['description']}")

        # System information
        st.markdown("### System Information")
        if torch.cuda.is_available():
            gpu_info = torch.cuda.get_device_properties(0)
            st.success(f"🎮 GPU: {gpu_info.name}")
            st.info(f"📝 VRAM: {gpu_info.total_memory/1024**3:.1f}GB")
        else:
            st.warning("⚠️ No GPU detected - Processing will be slower")
        
        # Usage statistics
        stats = st.session_state.manager.get_transcript_stats()
        if stats:
            st.markdown("### Usage Statistics")
            st.metric("Total Transcripts", stats['total_transcripts'])
            st.metric("Storage Used", format_size(stats['total_size']))

def main():
    # Initialize and check
    init_session_state()
    check_initialization()
    render_sidebar()
    
    # Main content
    st.markdown('<h1 class="main-title">Audio & Video Transcriber</h1>', unsafe_allow_html=True)
    
    # Input method selection
    input_method = st.radio(
        "Select Input Method",
        ["YouTube URL", "Audio File", "YouTube Playlist"],
        horizontal=True,
        help="Choose your input source"
    )
    
    if input_method == "YouTube URL":
        url = st.text_input(
            "Enter YouTube URL",
            help="Paste a YouTube video URL here"
        )
        
        if url:
            try:
                if not VideoUtility.validate_url(url):
                    st.error("Invalid YouTube URL. Please enter a valid video URL.")
                    st.stop()
                handle_single_video(url)
            except Exception as e:
                st.error(f"Error: {str(e)}")
                
    elif input_method == "Audio File":
        handle_audio_upload()
        
    else:  # YouTube Playlist
        url = st.text_input(
            "Enter YouTube Playlist URL",
            help="Paste a YouTube playlist URL here"
        )
        
        if url:
            try:
                if not is_playlist_url(url):
                    st.error("Invalid playlist URL. Please enter a valid YouTube playlist URL.")
                    st.stop()
                handle_playlist(url)
            except Exception as e:
                st.error(f"Error: {str(e)}")
    
    # Show recent transcriptions
    show_recent_transcriptions()

def handle_audio_upload():
    """Handle audio file upload and processing"""
    uploaded_file = st.file_uploader(
        "Upload Audio File",
        type=['mp3', 'wav', 'm4a', 'ogg', 'flac', 'aac', 'wma'],
        help="Supported formats: MP3, WAV, M4A, OGG, FLAC, AAC, WMA"
    )
    
    if uploaded_file:
        # Show file information
        file_info = {
            'name': uploaded_file.name,
            'size': uploaded_file.size,
            'type': uploaded_file.type,
            'extension': os.path.splitext(uploaded_file.name)[1].lower()
        }
        
        # File information display
        st.markdown("### File Information")
        cols = st.columns(3)
        
        with cols[0]:
            st.metric("File Size", format_size(file_info['size']))
        with cols[1]:
            st.metric("Format", file_info['extension'][1:].upper())
        with cols[2]:
            validation_result = AudioUtility.validate_audio_file(file_info)
            st.metric(
                "Status", 
                "✅ Valid" if validation_result['valid'] else "❌ Invalid",
                help=validation_result.get('message', '')
            )
        
        # Process if valid
        if validation_result['valid']:
            if st.button("Start Transcription", use_container_width=True):
                process_audio_file(uploaded_file, file_info)

def process_audio_file(uploaded_file, file_info):
    """Process uploaded audio file"""
    try:
        with st.spinner("Processing..."):
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            # Save uploaded file temporarily
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_info['extension']) as tmp_file:
                tmp_file.write(uploaded_file.getvalue())
                tmp_path = tmp_file.name
            
            try:
                # Process the audio
                doc_path = st.session_state.manager.transcribe_audio_file(
                    tmp_path,
                    file_info['name'],
                    st.session_state.model_key,
                    lambda p, s: (
                        progress_bar.progress(p),
                        status_text.text(s)
                    )
                )
                
                if doc_path:
                    st.success("✅ Transcription completed!")
                    
                    # Download options
                    col1, col2 = st.columns(2)
                    with col1:
                        with open(doc_path, 'rb') as f:
                            st.download_button(
                                "📄 Download Transcript",
                                f,
                                file_name=os.path.basename(doc_path),
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                    
                    with col2:
                        if st.button("📋 Copy to Clipboard", use_container_width=True):
                            from docx import Document
                            doc = Document(doc_path)
                            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                            st.code(text, language=None)
                            st.success("Text copied to clipboard!")
                            
            finally:
                # Cleanup
                try:
                    os.unlink(tmp_path)
                except:
                    pass
                    
    except Exception as e:
        st.error(f"Error processing audio file: {str(e)}")
    

def handle_playlist(url):
    """Handle playlist processing"""
    try:
        # Get playlist info
        playlist_info = extract_playlist_info(url)
        
        # Display playlist information
        st.markdown("### 📋 Playlist Information")
        
        # Basic info
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Total Videos", playlist_info['video_count'])
        with col2:
            st.metric("Total Duration", format_duration(playlist_info['total_duration']))
        with col3:
            model_info = ModelConfig.MODELS[st.session_state.model_option]
            est_time = ModelConfig.estimate_batch_processing_time(
                playlist_info['total_duration'],
                model_info['key']
            )
            st.metric("Est. Processing Time", format_duration(est_time))
        
        # Video list
        st.markdown("### 📼 Videos in Playlist")
        video_container = st.container()
        with video_container:
            for idx, video in enumerate(playlist_info['videos'], 1):
                with st.expander(f"{idx}. {video['title']}", expanded=False):
                    vcol1, vcol2 = st.columns([3, 1])
                    with vcol1:
                        st.markdown(f"**Duration:** {format_duration(video['duration'])}")
                        st.markdown(f"**Channel:** {video['channel']}")
                    with vcol2:
                        if video.get('thumbnail'):
                            st.image(video['thumbnail'], width=160)
        
        # Processing controls
        control_cols = st.columns(4)
        with control_cols[0]:
            process_all = st.button(
                "🚀 Process All",
                help="Process all videos in the playlist",
                use_container_width=True
            )
        with control_cols[1]:
            process_failed = st.button(
                "🔄 Retry Failed",
                help="Retry failed videos only",
                use_container_width=True
            )
        with control_cols[2]:
            resume_processing = st.button(
                "⏯️ Resume",
                help="Resume from previous progress",
                use_container_width=True
            )
        with control_cols[3]:
            stop_processing = st.button(
                "⏹️ Stop",
                help="Stop current processing",
                use_container_width=True
            )

        # Resume processing logic
        if resume_processing:
            resume_col1, resume_col2 = st.columns([3, 1])
            with resume_col1:
                resume_dir = st.text_input("Enter playlist directory path to resume from:")
            with resume_col2:
                confirm_resume = st.button("✅ Confirm Resume", use_container_width=True)
                
            if resume_dir and Path(resume_dir).exists() and confirm_resume:
                processor = PlaylistProcessor(st.session_state.manager)
                try:
                    result = processor.resume_playlist_processing(
                        Path(resume_dir),
                        st.session_state.model_key
                    )
                    if result:
                        display_processing_results(result)
                except Exception as e:
                    st.error(f"Error resuming processing: {str(e)}")

        # Regular processing logic
        if process_all or process_failed:
            try:
                # Initialize processor
                processor = PlaylistProcessor(st.session_state.manager)
                
                if stop_processing:
                    processor.stop_event.set()
                    st.warning("Stopping processing...")
                    st.stop()

                # Process playlist
                with st.spinner("Processing playlist..."):
                    result = processor.process_playlist(
                        url,
                        st.session_state.model_key,
                        retry_only=process_failed
                    )
                    
                    if result:
                        display_processing_results(result)
                        
            except Exception as e:
                st.error(f"Error processing playlist: {str(e)}")
    except Exception as e:
        st.error(f"Error loading playlist: {str(e)}")

def display_processing_results(result):
    """Display processing results and download options"""
    st.success(f"""✅ Playlist processing completed!
        - Successfully processed: {result['completed']} videos
        - Failed: {result['failed']} videos""")
    
    # Download options
    download_cols = st.columns(3)
    
    # Summary Report
    with download_cols[0]:
        if result.get('summary'):
            with open(result['summary'], 'rb') as f:
                st.download_button(
                    "📄 Download Summary Report",
                    f,
                    file_name=Path(result['summary']).name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
    
    # Error Log
    with download_cols[1]:
        error_log = Path(result['output_dir']) / "logs" / "error_log.txt"
        if error_log.exists():
            with open(error_log, 'rb') as f:
                st.download_button(
                    "⚠️ Download Error Log",
                    f,
                    file_name="error_log.txt",
                    mime="text/plain",
                    use_container_width=True
                )
    
    # All Transcripts
    with download_cols[2]:
        if result.get('output_dir') and Path(result['output_dir']).exists():
            try:
                # Create zip file of all transcripts
                create_download_zip(Path(result['output_dir']))
                
            except Exception as e:
                st.error(f"Error creating download package: {str(e)}")

def create_download_zip(output_dir: Path):
    """Create and offer download of transcripts zip file"""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.zip') as tmp_zip:
            with zipfile.ZipFile(tmp_zip.name, 'w') as zf:
                # Add transcripts
                for file in Path(output_dir).rglob('*.docx'):
                    if not file.name.startswith('00_playlist_summary'):
                        zf.write(file, f"transcripts/{file.name}")
                
                # Add summary
                summary_file = output_dir / "00_playlist_summary.docx"
                if summary_file.exists():
                    zf.write(summary_file, summary_file.name)
                
                # Add error log
                error_log = output_dir / "logs" / "error_log.txt"
                if error_log.exists():
                    zf.write(error_log, "logs/error_log.txt")

            # Offer download
            with open(tmp_zip.name, 'rb') as f:
                st.download_button(
                    "📦 Download All Files",
                    f,
                    file_name=f"transcripts_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip",
                    mime="application/zip",
                    use_container_width=True
                )
                
    except Exception as e:
        st.error(f"Error creating download package: {str(e)}")
    finally:
        # Cleanup
        try:
            if 'tmp_zip' in locals():
                Path(tmp_zip.name).unlink(missing_ok=True)
        except Exception:
            pass

def handle_single_video(url):
    """Handle single video processing"""
    try:
        # Show video preview
        video_info = st.session_state.preview.render_preview(
            url,
             st.session_state.model_key
        )
        
        if video_info:
            if st.button("Start Transcription", use_container_width=True):
                process_single_video(url, video_info)
    
    except Exception as e:
        st.error(f"Error loading video: {str(e)}")

def process_single_video(url, video_info):
    """Process a single video with progress tracking"""
    try:
        with st.spinner("Processing..."):
            # Create status containers
            progress_bar = st.progress(0)
            status_container = st.empty()
            error_container = st.empty()
            
            try:
                doc_path = st.session_state.manager.process_video(
                    url,
                    st.session_state.model_key,
                    lambda p, s: (
                        progress_bar.progress(p),
                        status_container.text(s)
                    )
                )
                
                if doc_path:
                    st.success("✅ Transcription completed!")
                    
                    # Download options
                    col1, col2 = st.columns(2)
                    with col1:
                        with open(doc_path, 'rb') as f:
                            st.download_button(
                                "📄 Download Transcript",
                                f,
                                file_name=os.path.basename(doc_path),
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                    
                    with col2:
                        # Option to copy to clipboard
                        if st.button("📋 Copy to Clipboard", use_container_width=True):
                            from docx import Document
                            doc = Document(doc_path)
                            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                            st.code(text, language=None)
                            st.success("Text copied to clipboard!")
                
            except Exception as e:
                error_container.error(f"Transcription failed: {str(e)}")
    
    except Exception as e:
        st.error(f"Error during transcription: {str(e)}")

def show_recent_transcriptions():
    """Show recent transcriptions section"""
    st.markdown("---")
    st.header("Recent Transcriptions")
    
    recent_files = st.session_state.manager.list_recent_transcripts()
    if recent_files:
        # Group files by date
        from itertools import groupby
        from datetime import datetime
        
        files_by_date = {}
        for file_path in recent_files:
            date = datetime.fromtimestamp(file_path.stat().st_mtime).strftime("%Y-%m-%d")
            if date not in files_by_date:
                files_by_date[date] = []
            files_by_date[date].append(file_path)
        
        # Display files by date
        for date in sorted(files_by_date.keys(), reverse=True):
            with st.expander(f"📅 {date}", expanded=(date == max(files_by_date.keys()))):
                for file_path in sorted(files_by_date[date], key=lambda x: x.stat().st_mtime, reverse=True):
                    col1, col2, col3 = st.columns([2, 1, 1])
                    
                    with col1:
                        st.markdown(f"**{file_path.stem}**")
                    
                    with col2:
                        modified_time = datetime.fromtimestamp(file_path.stat().st_mtime).strftime("%H:%M:%S")
                        st.text(f"⏰ {modified_time}")
                    
                    with col3:
                        with open(file_path, 'rb') as f:
                            st.download_button(
                                "📄 Download",
                                f,
                                file_name=file_path.name,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=str(file_path)
                            )
    else:
        st.info("No recent transcriptions found")

if __name__ == "__main__":
    main()