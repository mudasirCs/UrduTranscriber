import streamlit as st
import whisper
import yt_dlp
import os
import torch
import ffmpeg
from datetime import datetime
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import docx.oxml.shared
from docx.opc.constants import RELATIONSHIP_TYPE
from pathlib import Path
import shutil
from typing import Optional, Dict, Any
from urllib.error import HTTPError
import threading
from queue import Queue
import time
from .utils import AudioUtility, ModelConfig, VideoUtility, extract_playlist_info, format_duration, sanitize_filename, get_timestamp

class TranscriptionManager:
    def __init__(self):
        self.dirs = {
            'temp': Path("temp").absolute(),
            'output': Path("transcripts").absolute(),
            'cache': Path("cache").absolute()
        }
        # Create directories if they don't exist
        for dir_path in self.dirs.values():
            dir_path.mkdir(parents=True, exist_ok=True)
            
        # Initialize thread-safe progress tracking
        self.progress_queue = Queue()
        self.progress_lock = threading.Lock()
        self.download_progress = 0.0
        self.current_status = ""

    @staticmethod
    def add_hyperlink(paragraph, text, url):
        """Add a hyperlink to a paragraph"""
        try:
            part = paragraph.part
            r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
            
            hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
            hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)
            
            new_run = docx.oxml.shared.OxmlElement('w:r')
            rPr = docx.oxml.shared.OxmlElement('w:rPr')
            
            rStyle = docx.oxml.shared.OxmlElement('w:rStyle')
            rStyle.set(docx.oxml.shared.qn('w:val'), 'Hyperlink')
            rPr.append(rStyle)
            
            new_run.append(rPr)
            new_run.text = text
            hyperlink.append(new_run)
            paragraph._p.append(hyperlink)
        except Exception as e:
            raise Exception(f"Failed to add hyperlink: {str(e)}")

    def _download_progress_hook(self, d):
        """Progress hook for yt-dlp"""
        if d['status'] == 'downloading':
            try:
                total = d.get('total_bytes') or d.get('total_bytes_estimate', 0)
                downloaded = d.get('downloaded_bytes', 0)
                if total > 0:
                    with self.progress_lock:
                        self.download_progress = (downloaded / total)
                        
                # Update status with speed and ETA
                speed = d.get('speed', 0)
                if speed:
                    speed_mb = speed / 1024 / 1024  # Convert to MB/s
                    eta = d.get('eta', 0)
                    self.current_status = f"Downloading: {speed_mb:.1f} MB/s - ETA: {eta}s"
                    
            except Exception:
                pass
        elif d['status'] == 'error':
            self.current_status = f"Download error: {d.get('error', 'Unknown error')}"

    def _verify_audio_file(self, file_path: Path) -> bool:
        """Verify the downloaded audio file"""
        try:
            # Check file exists
            if not file_path.exists():
                return False
            
            # Check file size
            if file_path.stat().st_size < 1024:  # Less than 1KB
                return False
            
            # Try to open with ffmpeg to verify format
            try:
                probe = ffmpeg.probe(str(file_path))
                if not probe or 'streams' not in probe:
                    return False
                # Verify it has an audio stream
                has_audio = any(stream['codec_type'] == 'audio' for stream in probe['streams'])
                if not has_audio:
                    return False
            except ffmpeg.Error:
                return False
            
            return True
        except Exception as e:
            st.warning(f"File verification failed: {str(e)}")
            return False

    def download_audio(self, video_url: str, max_retries: int = 3, retry_delay: int = 5, 
                    output_path: Optional[str] = None) -> Dict[str, Any]:
        """Download audio from YouTube video with retries and progress tracking"""
        self.download_progress = 0.0
        self.current_status = "Starting download..."
        
        download_url = VideoUtility.get_clean_video_url(video_url)

        for attempt in range(max_retries):
            try:
                # Configure output path
                if output_path:
                    output_template = output_path
                    output_dir = Path(output_path).parent
                else:
                    output_dir = self.dirs['temp']
                    output_template = str(output_dir / '%(id)s.%(ext)s')
                
                # Ensure directory exists
                output_dir.mkdir(parents=True, exist_ok=True)

                # First try to get video info separately
                info_opts = {
                    'quiet': True,
                    'no_warnings': True,
                    'extract_flat': False,
                    'format': 'bestaudio/best',
                    'ignoreerrors': True,
                    'no_color': True
                }

                with yt_dlp.YoutubeDL(info_opts) as ydl_info:
                    try:
                        # Get video info first
                        video_info = ydl_info.extract_info(download_url, download=False)
                        if not video_info:
                            raise Exception("Could not fetch video information")
                            
                        # Check if video is available and has duration
                        if video_info.get('duration') is None:
                            # Try an alternative format if duration is not available
                            info_opts['format'] = 'bestaudio'
                            with yt_dlp.YoutubeDL(info_opts) as ydl_info2:
                                video_info = ydl_info2.extract_info(download_url, download=False)
                                if not video_info or video_info.get('duration') is None:
                                    raise Exception("Could not determine video duration")
                    except Exception as e:
                        if "Private video" in str(e):
                            raise Exception("This is a private video")
                        elif "Sign in" in str(e):
                            raise Exception("This video requires age verification or sign-in")
                        raise
                
                # Configure yt-dlp options for download                
                
                ydl_opts = {
                    'format': 'bestaudio[ext=m4a]/bestaudio/best',
                    'outtmpl': output_template,
                    'quiet': True,
                    'no_warnings': True,
                    'extract_audio': True,
                    'postprocessors': [{
                        'key': 'FFmpegExtractAudio',
                        'preferredcodec': 'mp3',
                        'preferredquality': '192'
                    }],
                    # Retry settings
                    'retries': 10,
                    'fragment_retries': 10,
                    'skip_unavailable_fragments': False,
                    'socket_timeout': 30,
                    # Rate limiting
                    'sleep_interval': 1,
                    'max_sleep_interval': 5,
                    'sleep_interval_requests': 1,
                    'extractor_args': {'youtube': {'skip': ['dash', 'hls']}},
                    # Additional options
                    'ignoreerrors': False,
                    'no_color': True,
                    'geo_bypass': True,
                    'geo_bypass_country': 'US',
                    'socket_timeout': 30,
                    'retry_sleep_functions': {
                        'http': lambda x: 5 * (x + 1),
                        'fragment': lambda x: 5 * (x + 1),
                        'file_access': lambda x: 5 * (x + 1),
                    },
                    'nocheckcertificate': True,
                    'prefer_ffmpeg': True,
                    'keepvideo': False,
                    'progress_hooks': [self._download_progress_hook]
                }
                
                # Download the audio
                with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                    info = ydl.extract_info(download_url, download=True)
                    video_id = info['id']
                    
                    # Check if video is available
                    if info.get('is_live', False):
                        raise Exception("Cannot download live streams")
                    
                    # Look for the MP3 file
                    mp3_file = None
                    check_attempts = 20
                    for _ in range(check_attempts):
                        if output_path and Path(output_path).exists():
                            mp3_file = Path(output_path)
                            break
                        else:
                            for file in output_dir.glob("*.mp3"):
                                if video_id in str(file):
                                    mp3_file = file
                                    break
                        if mp3_file:
                            break
                        time.sleep(0.5)
                    
                    if not mp3_file:
                        raise Exception(f"MP3 file not found after download. Video ID: {video_id}")
                    
                    # Verify file
                    if not self._verify_audio_file(mp3_file):
                        raise Exception("Audio file verification failed")
                    
                    # Return comprehensive info
                    return {
                        'path': str(mp3_file),
                        'video_id': info['id'],
                        'title': info.get('title', video_info.get('title', 'Unknown Title')),
                        'duration': info.get('duration', video_info.get('duration', 0)),
                        'channel': info.get('channel', video_info.get('channel', 'Unknown Channel')),
                        'upload_date': info.get('upload_date', video_info.get('upload_date', 'Unknown Date')),
                        'filesize': mp3_file.stat().st_size,
                        'format': 'mp3',
                        'bitrate': '192k',
                        'url': video_url
                    }
            
            except HTTPError as he:
                if he.code == 429:  # Too Many Requests
                    wait_time = retry_delay * (2 ** attempt)  # Exponential backoff
                    self.current_status = f"Rate limit detected. Waiting {wait_time} seconds..."
                    time.sleep(wait_time)
                    continue
                raise
            
            except Exception as e:
                error_msg = str(e)
                
                # Handle specific error cases
                if "Video unavailable" in error_msg:
                    raise Exception("Video is unavailable or private")
                elif "Copyright infringement" in error_msg:
                    raise Exception("Video not accessible due to copyright restrictions")
                elif "Sign in" in error_msg:
                    raise Exception("Video requires authentication")
                elif "This live event has ended" in error_msg:
                    raise Exception("This live stream has ended")
                
                if attempt < max_retries - 1:
                    wait_time = retry_delay * (attempt + 1)
                    if "429" in error_msg:
                        wait_time *= 2
                    self.current_status = f"Download failed. Retrying in {wait_time} seconds..."
                    time.sleep(wait_time)
                    continue
                else:
                    raise Exception(f"Failed to download audio after {max_retries} attempts: {error_msg}")
        
        raise Exception("Failed to download audio: Max retries exceeded")

    @staticmethod
    @st.cache_resource
    def load_model(model_name):
        """Load and cache Whisper model"""
        try:
            return whisper.load_model(model_name)
        except Exception as e:
            raise Exception(f"Failed to load model {model_name}: {str(e)}")

    def process_video(self, url: str, model_name: str = "base", progress_callback=None):
        """Process video and generate transcript with progress tracking"""
        audio_info = None
        try:
            if progress_callback:
                progress_callback(0.1, "Loading model...")

            # Add these CUDA optimization settings
            if torch.cuda.is_available():
                torch.cuda.empty_cache()
                torch.backends.cuda.matmul.allow_tf32 = True
                torch.backends.cudnn.allow_tf32 = True            

            model = TranscriptionManager.load_model(model_name)
            
            if progress_callback:
                progress_callback(0.2, "Downloading audio...")
            audio_info = self.download_audio(url, max_retries=3, retry_delay=5)
            
            if progress_callback:
                progress_callback(0.4, "Transcribing...")
            
            # Clear GPU memory before processing
            if torch.cuda.is_available():
                torch.cuda.empty_cache()
                model = model.cuda()
            
            model_info = None
            for display_name, config in ModelConfig.MODELS.items():
                if config['key'] == model_name:
                    model_info = {
                        'name': display_name,
                        'parameters': config['size'],
                        'key': model_name
                    }
                    break

            result = model.transcribe(
                audio_info['path'],
                language="ur",
                task="transcribe",
                fp16=False
            )
            
            if model_info:
                audio_info['model'] = model_info

            if progress_callback:
                progress_callback(0.8, "Generating document...")
            return self.save_transcript(result, audio_info, url)
            
        finally:
            if audio_info and 'path' in audio_info and os.path.exists(audio_info['path']):
                try:
                    os.remove(audio_info['path'])
                except:
                    pass

    def process_single_video(self, audio_path: str, video_url: str, video_title: str, 
                            progress_callback=None, metadata: Optional[Dict[str, Any]] = None,
                            model_name: str = "base"):  # Added model_name parameter
        """Process a single video from existing audio file"""
        try:
            # Load model
            if progress_callback:
                progress_callback(0.1, f"Loading {model_name} model...")

            model = self.load_model(model_name)
            
            # Clear GPU memory
            if torch.cuda.is_available():
                torch.cuda.empty_cache()
            
            if progress_callback:
                progress_callback(0.3, "Starting transcription...")
                
            # Transcribe
            result = model.transcribe(
                audio_path,
                language="ur",
                task="transcribe",
                fp16=False
            )
            
            if progress_callback:
                progress_callback(0.8, "Generating document...")
            
            # Use metadata if provided, otherwise create default info dict
            info = metadata or {
                'title': video_title,
                'path': audio_path,
                'duration': 0,
                'channel': 'Unknown'
            }

            for display_name, config in ModelConfig.MODELS.items():
                if config['key'] == model_name:
                    info['model'] = {
                        'name': display_name,  # Use the display name (e.g., "Large (1550M parameters)")
                        'parameters': config['size'],  # Get the parameters info
                        'key': model_name     # Keep the key (e.g., "large")
                    }
                    break
            else:
                # Fallback if model not found in config
                info['model'] = {
                    'name': model_name.title(),
                    'parameters': 'Unknown size',
                    'key': model_name
                }
            
            model_display_name = None
            model_params = None
            for display_name, config in ModelConfig.MODELS.items():
                if config['key'] == model_name:
                    model_display_name = display_name
                    model_params = config['size']
                    break

            if not model_display_name:
                # Fallback for model names
                model_map = {
                    'tiny': 'Tiny (fastest)',
                    'base': 'Base',
                    'small': 'Small',
                    'medium': 'Medium',
                    'large': 'Large'
                }
                model_display_name = model_map.get(model_name, model_name.title())
                model_params = ModelConfig.MODELS.get(model_display_name, {}).get('size', '')
            
            info['model'] = {
                'name': model_name,
                'display_name': model_display_name,
                'parameters': model_params
            }

            # Save transcript
            if progress_callback:
                progress_callback(0.9, "Saving transcript...")
                
            transcript_path = self.save_transcript(result, info, video_url)
            
            if progress_callback:
                progress_callback(1.0, "Complete!")
                
            return transcript_path
            
        except Exception as e:
            if progress_callback:
                progress_callback(0, f"Error: {str(e)}")
            raise Exception(f"Failed to process video: {str(e)}")

    
    def save_transcript(self, result: Dict[str, Any], info: Dict[str, Any], url: str, 
                    output_path: Optional[str] = None):
        """Save transcript to Word document with optional output path"""
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading(info['title'], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Video Information
            doc.add_heading('Video Information', level=1)
            p = doc.add_paragraph('Video URL: ')
            self.add_hyperlink(p, url, url)
            
            # Info table
            info_table = doc.add_table(rows=4, cols=2)
            info_table.style = 'Table Grid'
            
            # Format model info
            model_info = info.get('model', {})
            if model_info:
                model_text = f"{model_info['name']}"
                if model_info.get('parameters'):
                    model_text += f" ({model_info['parameters']})"
            else:
                model_text = "Unknown Model"
            
            # Add information to table
            rows = info_table.rows
            
            # Model information
            rows[0].cells[0].text = 'Model Used'
            rows[0].cells[1].text = model_text
            
            # Channel information
            rows[1].cells[0].text = 'Channel'
            rows[1].cells[1].text = info.get('channel', 'Unknown')
            
            # Duration information
            rows[2].cells[0].text = 'Duration'
            rows[2].cells[1].text = format_duration(info.get('duration', 0))
            
            # Processing date
            rows[3].cells[0].text = 'Processing Date'
            rows[3].cells[1].text = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            
            # Rest of the transcript generation...
            # Complete transcript
            doc.add_heading('Complete Transcript', level=1)
            p = doc.add_paragraph()
            text_run = p.add_run(result['text'])
            text_run.font.name = 'Jameel Noori Nastaleeq'
            text_run.font.size = Pt(14)
            
            # Add page break
            doc.add_page_break()
            
            # Timestamped transcript
            doc.add_heading('Timestamped Transcript', level=1)
            
            # Extract video ID
            video_id = None
            if 'v=' in url:
                video_id = url.split('v=')[1].split('&')[0]
            elif 'youtu.be/' in url:
                video_id = url.split('youtu.be/')[1].split('?')[0]
            else:
                video_id = url.split('/')[-1].split('?')[0]
            
            for segment in result['segments']:
                p = doc.add_paragraph()
                timestamp_seconds = int(segment['start'])
                timestamp_url = f"https://youtube.com/watch?v={video_id}&t={timestamp_seconds}"
                self.add_hyperlink(
                    p, 
                    f"[{format_duration(segment['start'])} - {format_duration(segment['end'])}]",
                    timestamp_url
                )
                
                text_run = p.add_run('\n' + segment['text'].strip())
                text_run.font.name = 'Jameel Noori Nastaleeq'
                text_run.font.size = Pt(14)
            
            # Save document
            if output_path:
                doc_path = Path(output_path)
            else:
                timestamp = get_timestamp()
                safe_title = sanitize_filename(info['title'])
                doc_path = self.dirs['output'] / f"{timestamp}_{safe_title}.docx"
            
            doc_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(doc_path))
            return str(doc_path)
            
        except Exception as e:
            raise Exception(f"Failed to save transcript: {str(e)}")
        

    def list_recent_transcripts(self, limit=5):
        """Get list of recent transcriptions with recursive search"""
        try:
            # Get all .docx files from output directory and its subdirectories
            all_files = []
            for file_path in self.dirs['output'].rglob('*.docx'):
                if file_path.is_file() and not file_path.name.startswith('00_playlist_summary'):
                    all_files.append(file_path)
            
            # Sort by modification time and return most recent
            recent_files = sorted(
                all_files,
                key=lambda x: x.stat().st_mtime,
                reverse=True
            )[:limit]
            
            return recent_files
        except Exception as e:
            st.warning(f"Error listing recent transcripts: {str(e)}")
            return []

    def cleanup_temp_files(self):
        """Clean up temporary files"""
        try:
            # Clean temp directory
            for file in self.dirs['temp'].glob("*"):
                try:
                    if file.is_file():
                        file.unlink()
                except Exception as e:
                    st.warning(f"Failed to delete temporary file {file}: {str(e)}")
            
            # Clean cache directory
            for file in self.dirs['cache'].glob("*"):
                try:
                    if file.is_file() and (datetime.now() - datetime.fromtimestamp(file.stat().st_mtime)).days > 7:
                        file.unlink()
                except Exception as e:
                    st.warning(f"Failed to delete cache file {file}: {str(e)}")
                    
        except Exception as e:
            st.warning(f"Error cleaning up temporary files: {str(e)}")

    def get_download_progress(self) -> float:
        """Get current download progress"""
        with self.progress_lock:
            return self.download_progress

    def get_current_status(self) -> str:
        """Get current status message"""
        return self.current_status

    def get_transcript_stats(self):
        """Get statistics about transcripts"""
        try:
            stats = {
                'total_transcripts': 0,
                'total_playlists': 0,
                'total_duration': 0,
                'total_size': 0
            }
            
            # Count regular transcripts
            for file in self.dirs['output'].rglob('*.docx'):
                if file.name.startswith('00_playlist_summary'):
                    stats['total_playlists'] += 1
                else:
                    stats['total_transcripts'] += 1
                stats['total_size'] += file.stat().st_size
            
            return stats
        except Exception as e:
            st.warning(f"Error getting transcript stats: {str(e)}")
            return None

    def verify_system_requirements(self):
        """Verify system requirements and configurations"""
        requirements_met = True
        issues = []
        
        # Check FFmpeg
        try:
            ffmpeg.probe(None)
        except ffmpeg.Error:
            requirements_met = False
            issues.append("FFmpeg not found. Please install FFmpeg.")
        except Exception:
            pass
        
        # Check GPU availability
        if torch.cuda.is_available():
            gpu_info = f"GPU available: {torch.cuda.get_device_name(0)}"
            gpu_memory = torch.cuda.get_device_properties(0).total_memory / 1024**3
            if gpu_memory < 4:
                issues.append(f"Low GPU memory: {gpu_memory:.1f}GB. Some models may not work.")
        else:
            requirements_met = False
            issues.append("No GPU found. Processing will be slower.")
        
        # Check directories
        for dir_name, dir_path in self.dirs.items():
            if not dir_path.exists():
                try:
                    dir_path.mkdir(parents=True)
                except Exception:
                    requirements_met = False
                    issues.append(f"Cannot create {dir_name} directory.")
        
        # Check write permissions
        for dir_path in self.dirs.values():
            try:
                test_file = dir_path / "test.txt"
                test_file.touch()
                test_file.unlink()
            except Exception:
                requirements_met = False
                issues.append(f"No write permission in {dir_path}")
        
        return requirements_met, issues

    def estimate_processing_time(self, duration: int, model_name: str) -> int:
        """Estimate processing time in seconds"""
        # Base processing speed factors
        speed_factors = {
            'tiny': 0.5,
            'base': 1.0,
            'small': 2.0,
            'medium': 3.0,
            'large': 4.0
        }
        
        # Get speed factor for model
        speed_factor = speed_factors.get(model_name, 1.0)
        
        # Calculate processing time
        processing_time = duration * speed_factor
        
        # Add overhead for file operations
        overhead = 30  # Base overhead in seconds
        if duration > 3600:  # Additional overhead for long videos
            overhead += duration * 0.05
        
        return int(processing_time + overhead)

    def transcribe_audio_file(self, audio_path, filename, model_name="base", progress_callback=None):
        """Process uploaded audio file"""
        try:
            if progress_callback:
                progress_callback(0.1, "Loading model...")
            model = self.load_model(model_name)
            
            if progress_callback:
                progress_callback(0.3, "Preparing audio...")
                
            # Convert to MP3 if needed
            if not audio_path.endswith('.mp3'):
                temp_mp3 = os.path.join(self.dirs['temp'], 'temp.mp3')
                if not AudioUtility.convert_to_mp3(audio_path, temp_mp3):
                    raise Exception("Failed to convert audio format")
                audio_path = temp_mp3
            
            if progress_callback:
                progress_callback(0.4, "Transcribing...")
                
            # Get duration for document
            duration = AudioUtility.get_audio_duration(audio_path)
            
            # Transcribe
            result = model.transcribe(
                audio_path,
                language="ur",
                task="transcribe",
                fp16=False
            )
            
            if progress_callback:
                progress_callback(0.8, "Generating document...")
            
            # Create info dict for document
            info = {
                'title': os.path.splitext(filename)[0],
                'duration': duration or 0,
                'channel': 'Local Audio',
                'upload_date': datetime.now().strftime('%Y-%m-%d')
            }
            
            # Generate document
            doc_path = self.save_transcript(result, info, "Local Audio File")
            
            if progress_callback:
                progress_callback(1.0, "Complete!")
                
            return doc_path
            
        except Exception as e:
            raise Exception(f"Error processing audio: {str(e)}")
            
        finally:
            # Cleanup temporary files
            if 'temp_mp3' in locals() and os.path.exists(temp_mp3):
                try:
                    os.remove(temp_mp3)
                except:
                    pass

    def get_transcript_stats(self):
        """Get statistics about transcriptions"""
        try:
            transcripts = list(Path(self.dirs['output']).glob('*.docx'))
            total_size = sum(f.stat().st_size for f in transcripts)
            
            return {
                'total_transcripts': len(transcripts),
                'total_size': total_size,
                'latest_transcript': max(transcripts, key=lambda x: x.stat().st_mtime).name if transcripts else None
            }
        except Exception:
            return None
