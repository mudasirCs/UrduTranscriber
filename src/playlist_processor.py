import concurrent.futures
import json
import threading
import time
import pandas as pd
import streamlit as st
from pathlib import Path
from dataclasses import asdict, dataclass
from typing import Any, Dict, Optional, List
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import whisper
import yt_dlp
from .utils import ModelConfig, extract_playlist_info, format_duration, get_timestamp
import re
import torch
import shutil
from urllib.error import HTTPError

@dataclass
class VideoTask:
    idx: int
    url: str
    title: str
    duration: int
    status: str = 'pending'  # 'pending', 'downloading', 'transcribing', 'completed', 'failed'
    error: Optional[str] = None
    transcript_path: Optional[str] = None
    audio_path: Optional[str] = None
    last_attempt: Optional[str] = None
    processing_progress: float = 0
    stop_requested: bool = False
    attempt_count: int = 0
    download_size: int = 0
    downloaded_size: int = 0
    download_speed: float = 0
    stop_requested: bool = False
    channel: str = 'Unknown' 
    video_id: Optional[str] = None
    thumbnail: Optional[str] = None
    view_count: int = 0
    upload_date: str = 'Unknown'
    processed_duration: float = 0   
    

    def to_dict(self):
            """Convert task to dictionary for saving progress"""
            return {
                'idx': self.idx,
                'url': self.url,
                'title': self.title,
                'duration': self.duration,
                'status': self.status,
                'error': self.error,
                'transcript_path': str(self.transcript_path) if self.transcript_path else None,
                'audio_path': str(self.audio_path) if self.audio_path else None,
                'last_attempt': self.last_attempt,
                'attempt_count': self.attempt_count,
                'processing_progress': self.processing_progress,
                'channel': self.channel,
                'video_id': self.video_id,
                'view_count': self.view_count,
                'upload_date': self.upload_date,
                'processed_duration': self.processed_duration
            }

    @classmethod
    def from_dict(cls, data):
        """Create task from dictionary when loading progress"""
        return cls(**data)

    def save_progress(self):
        """Save current progress to file"""
        if self.current_playlist_dir:
            progress_file = self.current_playlist_dir / "progress.json"
            try:
                # Create progress data
                progress_data = {
                    'tasks': {
                        vid_id: task.to_dict() 
                        for vid_id, task in self.tasks.items()
                    },
                    'total_duration': self.total_duration,
                    'timestamp': datetime.now().isoformat(),
                    'model_name': getattr(self, 'model_name', 'unknown'),
                    'completed_count': len([t for t in self.tasks.values() if t.status == 'completed']),
                    'failed_count': len([t for t in self.tasks.values() if t.status == 'failed']),
                    'total_processed_duration': sum(t.processed_duration for t in self.tasks.values())
                }

                # Save to temporary file first
                temp_file = progress_file.with_suffix('.tmp')
                with open(temp_file, 'w', encoding='utf-8') as f:
                    json.dump(progress_data, f, indent=2, ensure_ascii=False)

                # Rename temporary file to actual progress file
                temp_file.replace(progress_file)
                
            except Exception as e:
                st.warning(f"Failed to save progress: {str(e)}")

    def load_progress(self, playlist_dir: Path) -> bool:
        """Load progress from file"""
        progress_file = playlist_dir / "progress.json"
        if progress_file.exists():
            try:
                with open(progress_file, 'r', encoding='utf-8') as f:
                    progress_data = json.load(f)

                # Restore tasks
                self.tasks = {
                    vid_id: VideoTask.from_dict(task_data)
                    for vid_id, task_data in progress_data['tasks'].items()
                }

                # Restore other progress information
                self.total_duration = progress_data.get('total_duration', 0)
                self.model_name = progress_data.get('model_name', 'unknown')
                
                # Calculate progress
                completed_duration = progress_data.get('total_processed_duration', 0)
                if self.total_duration > 0:
                    progress = completed_duration / self.total_duration
                    st.info(f"Restored progress: {progress:.1%} complete")
                
                return True
                
            except Exception as e:
                st.warning(f"Failed to load progress: {str(e)}")
            
        return False

    def update_task_status(self, video_id: str, status: str, error: Optional[str] = None, 
                        transcript_path: Optional[str] = None):
        """Update task status with thread safety"""
        with self.lock:
            if video_id in self.tasks:
                task = self.tasks[video_id]
                task.status = status
                task.last_attempt = datetime.now().isoformat()
                
                if error:
                    task.error = error
                if transcript_path:
                    task.transcript_path = transcript_path
                
                # Save progress after status update
                self.save_progress()

    def get_task_statistics(self):
        """Get comprehensive processing statistics"""
        with self.lock:
            stats = {
                'total': len(self.tasks),
                'completed': len([t for t in self.tasks.values() if t.status == 'completed']),
                'failed': len([t for t in self.tasks.values() if t.status == 'failed']),
                'processing': len([t for t in self.tasks.values() 
                                if t.status in ['processing', 'downloading', 'transcribing']]),
                'pending': len([t for t in self.tasks.values() if t.status == 'pending']),
                'total_duration': sum(task.duration for task in self.tasks.values()),
                'processed_duration': sum(task.processed_duration for task in self.tasks.values())
            }
            return stats

    def cleanup_temporary_files(self):
        """Clean up temporary files"""
        try:
            if self.current_playlist_dir:
                audio_dir = self.current_playlist_dir / "audio"
                if audio_dir.exists():
                    for file in audio_dir.glob("*"):
                        try:
                            if file.is_file():
                                file.unlink()
                        except Exception as e:
                            st.warning(f"Failed to delete temporary file {file}: {str(e)}")
        except Exception as e:
            st.warning(f"Error cleaning up temporary files: {str(e)}")

    def resume_playlist_processing(self, playlist_dir: Path, model_name: str):
        """Resume playlist processing from saved progress"""
        if self.load_progress(playlist_dir):
            st.info("Resuming playlist processing...")
            
            # Get list of incomplete videos
            incomplete_videos = [
                vid_id for vid_id, task in self.tasks.items()
                if task.status in ['pending', 'failed', 'processing']
            ]
            
            if incomplete_videos:
                return self.process_playlist(None, model_name, incomplete_videos)
            else:
                st.success("All videos have been processed!")
                return None
        else:
            st.error("No progress file found to resume from")
            return None


class PlaylistProcessor:
    def __init__(self, transcription_manager=None):
            """Initialize PlaylistProcessor with additional attributes"""
            self.manager = transcription_manager
            self.tasks: Dict[str, VideoTask] = {}
            self.lock = threading.Lock()
            self.current_playlist_dir = None
            self.stop_event = threading.Event()
            self.model = None
            self.total_duration = 0
            self.processed_duration = 0
            self.model_name = None
            self.progress_file = None
            self.error_log_file = None

    def setup_playlist_directory(self, playlist_info):
        """Setup directory structure with error logging"""
        timestamp = get_timestamp()
        safe_title = self.sanitize_text(playlist_info['title']).replace(" ", "_")[:50]
        self.current_playlist_dir = self.manager.dirs['output'] / f"playlist_{timestamp}_{safe_title}"
        self.current_playlist_dir.mkdir(parents=True, exist_ok=True)
        
        # Create subdirectories
        (self.current_playlist_dir / "audio").mkdir(exist_ok=True)
        (self.current_playlist_dir / "transcripts").mkdir(exist_ok=True)
        (self.current_playlist_dir / "failed").mkdir(exist_ok=True)
        (self.current_playlist_dir / "logs").mkdir(exist_ok=True)
        
        # Setup log files
        self.error_log_file = self.current_playlist_dir / "logs" / "error_log.txt"
        self.progress_file = self.current_playlist_dir / "progress.json"
        
        return self.current_playlist_dir

    def log_error(self, video_id: str, error: str):
        """Log errors with timestamps"""
        if self.error_log_file:
            try:
                with open(self.error_log_file, 'a', encoding='utf-8') as f:
                    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    task = self.tasks.get(video_id)
                    title = task.title if task else 'Unknown'
                    f.write(f"[{timestamp}] Video: {title} ({video_id})\nError: {error}\n\n")
            except Exception as e:
                st.warning(f"Failed to log error: {str(e)}")

    def get_error_summary(self):
        """Get comprehensive error summary"""
        error_counts = {}
        error_details = []
        
        for task in self.tasks.values():
            if task.error:
                error_type = task.error.split(':')[0].strip()
                error_counts[error_type] = error_counts.get(error_type, 0) + 1
                error_details.append({
                    'video_id': task.video_id,
                    'title': task.title,
                    'error': task.error,
                    'attempts': task.attempt_count
                })
        
        return {
            'counts': dict(sorted(error_counts.items(), key=lambda x: x[1], reverse=True)),
            'details': error_details
        }

    def create_summary(self, playlist_info, output_dir: Path, model_name: str):
        """Create comprehensive summary with error details"""
        doc = Document()
        
        # Title
        title = doc.add_heading(f"Playlist: {playlist_info['title']}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Summary info
        doc.add_heading('Playlist Information', level=1)
        summary_table = doc.add_table(rows=8, cols=2)  # Added row for model info
        summary_table.style = 'Table Grid'
        
        # Get statistics and error summary
        stats = self.get_task_statistics()
        error_summary = self.get_error_summary()
        
        # Get model display info
        model_display_name = None
        model_params = None
        for display_name, config in ModelConfig.MODELS.items():
            if config['key'] == model_name:
                model_display_name = display_name
                model_params = config['size']
                break
        
        # Fill summary table
        rows = summary_table.rows
        rows[0].cells[0].text = 'Model Used'
        rows[0].cells[1].text = f"{model_display_name} ({model_params})" if model_display_name else model_name
        
        rows[1].cells[0].text = 'Total Videos'
        rows[1].cells[1].text = str(stats['total'])
        
        rows[2].cells[0].text = 'Successfully Processed'
        rows[2].cells[1].text = str(stats['completed'])
        
        rows[3].cells[0].text = 'Failed Videos'
        rows[3].cells[1].text = str(stats['failed'])
        
        rows[4].cells[0].text = 'Total Duration'
        rows[4].cells[1].text = format_duration(stats['total_duration'])
        
        rows[5].cells[0].text = 'Processed Duration'
        rows[5].cells[1].text = format_duration(stats['processed_duration'])
        
        rows[6].cells[0].text = 'Success Rate'
        success_rate = (stats['completed'] / stats['total'] * 100) if stats['total'] > 0 else 0
        rows[6].cells[1].text = f"{success_rate:.1f}%"
        
        rows[7].cells[0].text = 'Processing Date'
        rows[7].cells[1].text = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # Add error summary
        if error_summary['counts']:
            doc.add_heading('Error Summary', level=1)
            error_table = doc.add_table(rows=len(error_summary['counts']), cols=2)
            error_table.style = 'Table Grid'
            
            for idx, (error_type, count) in enumerate(error_summary['counts'].items()):
                row = error_table.rows[idx]
                row.cells[0].text = error_type
                row.cells[1].text = str(count)
        
        # Add detailed error log
        if error_summary['details']:
            doc.add_heading('Error Details', level=1)
            for error in error_summary['details']:
                p = doc.add_paragraph()
                p.add_run(f"Video: {error['title']}\n").bold = True
                p.add_run(f"Error: {error['error']}\n")
                p.add_run(f"Attempts: {error['attempts']}")
                doc.add_paragraph()  # Add spacing
        
        # Save summary
        summary_path = output_dir / "00_playlist_summary.docx"
        doc.save(str(summary_path))
        return summary_path

    @staticmethod
    def get_readable_error(error: str) -> str:
        """Convert error messages to user-friendly format"""
        common_errors = {
            'HTTP Error 429': 'Too many requests - Rate limit exceeded',
            'Video unavailable': 'Video is private or has been removed',
            'Sign in': 'Video requires age verification or sign-in',
            'copyright': 'Video not available due to copyright restrictions',
            'ffmpeg': 'Audio processing error',
            'cuda': 'GPU memory error',
            'out of memory': 'Out of memory error - Try a smaller model'
        }
        
        for key, readable in common_errors.items():
            if key.lower() in error.lower():
                return readable
        return error

    def save_progress(self):
        """Save current progress to file"""
        if self.current_playlist_dir:
            progress_file = self.current_playlist_dir / "progress.json"
            try:
                # Create progress data
                progress_data = {
                    'tasks': {
                        vid_id: task.to_dict() 
                        for vid_id, task in self.tasks.items()
                    },
                    'total_duration': self.total_duration,
                    'processed_duration': self.processed_duration,  # Include this in save
                    'timestamp': datetime.now().isoformat(),
                    'model_name': getattr(self, 'model_name', 'unknown')
                }

                # Save to temporary file first
                temp_file = progress_file.with_suffix('.tmp')
                with open(temp_file, 'w', encoding='utf-8') as f:
                    json.dump(progress_data, f, indent=2, ensure_ascii=False)

                # Rename temporary file to actual progress file
                temp_file.replace(progress_file)
                
            except Exception as e:
                st.warning(f"Failed to save progress: {str(e)}")

 
    def load_progress(self, playlist_dir: Path) -> bool:
        """Load progress from file"""
        progress_file = playlist_dir / "progress.json"
        if progress_file.exists():
            try:
                with open(progress_file, 'r') as f:
                    progress_data = json.load(f)
                self.tasks = {
                    vid_id: VideoTask.from_dict(task_data)
                    for vid_id, task_data in progress_data['tasks'].items()
                }
                self.total_duration = progress_data['total_duration']
                self.processed_duration = progress_data['processed_duration']
                return True
            except Exception as e:
                st.warning(f"Failed to load progress: {str(e)}")
        return False
 

    def setup_playlist_directory(self, playlist_info):
        """Create and setup playlist directory structure"""
        timestamp = get_timestamp()
        safe_title = self.sanitize_text(playlist_info['title']).replace(" ", "_")[:50]
        base_dir = self.manager.dirs['output'] / f"playlist_{timestamp}_{safe_title}"
        
        # Create directory structure
        base_dir.mkdir(parents=True, exist_ok=True)
        (base_dir / "audio").mkdir(exist_ok=True)
        (base_dir / "transcripts").mkdir(exist_ok=True)
        (base_dir / "failed").mkdir(exist_ok=True)
        
        self.current_playlist_dir = base_dir
        return base_dir

    def get_audio_path(self, video_id):
        """Get path for audio file"""
        return self.current_playlist_dir / "audio" / f"{video_id}.mp3"

    def get_transcript_path(self, video_id, title):
        """Get path for transcript file"""
        safe_title = self.sanitize_text(title).replace(" ", "_")[:50]
        return self.current_playlist_dir / "transcripts" / f"{video_id}_{safe_title}.docx"

    @staticmethod
    def sanitize_text(text: str) -> str:
        """Sanitize text to be XML compatible and safe for filenames"""
        if not text:
            return ""
        # Remove control characters but keep newlines and tabs
        text = ''.join(char for char in text if char in '\n\t' or (ord(char) >= 32 and ord(char) != 127))
        # Replace any remaining invalid characters
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        # Replace filesystem unsafe characters
        text = re.sub(r'[<>:"/\\|?*]', '_', text)
        # Remove leading/trailing spaces and dots
        text = text.strip('. ')
        return text

    def process_playlist(self, playlist_url: str, model_name: str = "base", retry_videos: List[str] = None, retry_only: bool = False):
        """Process playlist with optimized model loading and progress saving"""
        try:
            if not self.manager:
                raise Exception("Transcription manager not set")

            # Get playlist info if not retrying specific videos
            if not retry_videos and not retry_only:
                try:
                    playlist_info = extract_playlist_info(playlist_url)
                    if not playlist_info or 'videos' not in playlist_info:
                        raise Exception("Invalid playlist information")
                    
                    videos = playlist_info.get('videos', [])
                    if not videos:
                        raise Exception("No valid videos found in playlist")
                    
                    # Create directory structure
                    timestamp = get_timestamp()
                    safe_title = self.sanitize_text(playlist_info.get('title', 'Unknown')).replace(" ", "_")[:50]
                    self.current_playlist_dir = self.manager.dirs['output'] / f"playlist_{timestamp}_{safe_title}"
                    self.current_playlist_dir.mkdir(parents=True, exist_ok=True)
                    (self.current_playlist_dir / "audio").mkdir(exist_ok=True)
                    (self.current_playlist_dir / "transcripts").mkdir(exist_ok=True)
                    (self.current_playlist_dir / "failed").mkdir(exist_ok=True)
                    (self.current_playlist_dir / "logs").mkdir(exist_ok=True)

                    # Create tasks for each video
                    self.tasks.clear()
                    self.total_duration = 0
                    self.processed_duration = 0
                    valid_videos = 0

                    for idx, video in enumerate(videos, 1):
                        try:
                            # Validate video data
                            if not isinstance(video, dict):
                                st.warning(f"Skipping video {idx}: Invalid video data")
                                continue

                            # Get URL with validation
                            video_url = video.get('url', '')
                            if not video_url:
                                st.warning(f"Skipping video {idx}: Missing URL")
                                continue

                            # Extract video ID safely
                            video_id = None
                            try:
                                if 'v=' in video_url:
                                    video_id = video_url.split('v=')[1].split('&')[0]
                                elif '/watch/' in video_url:
                                    video_id = video_url.split('/watch/')[1].split('?')[0]
                                elif 'youtu.be/' in video_url:
                                    video_id = video_url.split('youtu.be/')[1].split('?')[0]
                                else:
                                    st.warning(f"Skipping video {idx}: Invalid URL format")
                                    continue
                            except Exception:
                                st.warning(f"Skipping video {idx}: Could not extract video ID")
                                continue

                            if not video_id or len(video_id) != 11:
                                st.warning(f"Skipping video {idx}: Invalid video ID")
                                continue

                            # Get duration safely
                            duration = video.get('duration', 0)
                            if duration is None:
                                duration = 0

                            # Create task
                            self.tasks[video_id] = VideoTask(
                                idx=idx,
                                url=video_url,
                                title=video.get('title', f'Video {idx}'),
                                duration=duration,
                                channel=video.get('channel', 'Unknown'),
                                video_id=video_id,
                                thumbnail=video.get('thumbnail'),
                                view_count=video.get('view_count', 0),
                                upload_date=video.get('upload_date', 'Unknown')
                            )

                            self.total_duration += duration
                            valid_videos += 1

                        except Exception as e:
                            st.warning(f"Error processing video {idx}: {str(e)}")
                            continue

                    if valid_videos == 0:
                        raise Exception("No valid videos found in playlist after processing")

                    st.info(f"Found {valid_videos} valid videos to process")

                except Exception as e:
                    raise Exception(f"Error processing playlist information: {str(e)}")


            else:
                # Handle retry logic
                if retry_only:
                    failed_videos = [vid_id for vid_id, task in self.tasks.items() 
                                if task.status == 'failed']
                    retry_videos = failed_videos
                
                if not retry_videos:
                    st.warning("No failed videos to retry")
                    return None
                    
                playlist_info = {'title': 'Retry Processing'}
                total_videos = len(retry_videos)
                
                # Reset status for retry videos
                for video_id in retry_videos:
                    if video_id in self.tasks:
                        task = self.tasks[video_id]
                        task.status = 'pending'
                        task.error = None
                        task.last_attempt = None
                        task.attempt_count = 0
                        task.processing_progress = 0

            # Load model once for all videos
            try:
                st.info(f"Loading {model_name} model...")
                if torch.cuda.is_available():
                    torch.cuda.empty_cache()
                self.model = whisper.load_model(model_name)
            except Exception as e:
                raise Exception(f"Failed to load model: {str(e)}")

            # Get the list of videos to process
            videos_to_process = retry_videos if retry_videos else list(self.tasks.keys())
            total_videos = len(videos_to_process)

            if total_videos == 0:
                st.warning("No videos to process")
                return None

            # Create status displays
            progress_bar = st.progress(0)
            status_text = st.empty()
            status_table = st.empty()
            
            # Process videos
            completed = 0
            failed = 0
            processed_duration = 0
            
            try:
                for idx, video_id in enumerate(videos_to_process, 1):
                    if self.stop_event.is_set():
                        st.warning("Processing stopped by user")
                        self.save_progress()
                        break

                    task = self.tasks.get(video_id)
                    if not task:
                        continue
                    
                    try:
                        # Update status display
                        self.display_status_table(status_table)
                        
                        # Process video
                        st.info(f"Processing video {idx}/{total_videos}: {task.title}")
                        result = self.process_single_video(video_id, task, model_name)
                        
                        if result and result.get('success'):
                            completed += 1
                            processed_duration += task.duration
                            task.processed_duration = task.duration
                            # Update overall progress based on duration
                            if self.total_duration > 0:
                                total_progress = processed_duration / self.total_duration
                                progress_bar.progress(total_progress)
                            st.success(f"Successfully processed: {task.title}")
                        else:
                            failed += 1
                            st.error(f"Failed to process: {task.title}")
                        
                        # Save progress after each video
                        self.save_progress()
                        
                    except Exception as e:
                        failed += 1
                        st.error(f"Error processing video {task.title}: {str(e)}")
                        self.save_progress()
                    
                    # Update status
                    status_text.text(f"Processed: {idx}/{total_videos} videos")
                    self.display_status_table(status_table)

                # Create summary if any videos were processed
                if completed > 0 or failed > 0:
                    # Process output directory
                    output_dir = self.current_playlist_dir
                    
                    # Move successful transcripts
                    successful_tasks = [task for task in self.tasks.values() 
                                    if task.status == 'completed' and task.transcript_path]
                    
                    for task in successful_tasks:
                        if task.transcript_path:
                            source = Path(task.transcript_path)
                            if source.exists():
                                dest = output_dir / source.name
                                shutil.move(str(source), str(dest))
                                task.transcript_path = str(dest)
                    
                    # Create and save summary
                    summary_path = self.create_summary(playlist_info, output_dir, model_name)
                    
                    return {
                        'summary': summary_path,
                        'completed': completed,
                        'failed': failed,
                        'output_dir': output_dir
                    }
                else:
                    st.warning("No videos were processed successfully")
                    return None

            finally:
                # Cleanup model
                if hasattr(self, 'model'):
                    del self.model
                    if torch.cuda.is_available():
                        torch.cuda.empty_cache()
                    
        except Exception as e:
            st.error(f"Error processing playlist: {str(e)}")
            self.save_progress()
            raise
        
    def process_single_video(self, video_id: str, task: VideoTask, model_name: str):
        """Process a single video using shared model instance"""
        max_retries = 5
        retry_delay = 10
        video_info = None
        
        for attempt in range(max_retries):
            try:
                self.update_task_status(video_id, 'processing')
                task.last_attempt = datetime.now().isoformat()
                task.attempt_count += 1
                
                # Download and process
                audio_info = None
                try:
                    # Get detailed video info first
                    ydl_opts = {
                        'quiet': True,
                        'no_warnings': True,
                        'extract_flat': False,
                        'format': 'best',
                        'ignoreerrors': True,
                        'no_color': True,
                        'socket_timeout': 30
                    }
                    
                    # Try to get video info
                    try:
                        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                            video_info = ydl.extract_info(task.url, download=False)
                            
                            if video_info:
                                # Update task with accurate information
                                task.title = video_info.get('title', task.title)
                                task.duration = video_info.get('duration', task.duration)
                                task.channel = video_info.get('channel', video_info.get('uploader', task.channel))
                                task.view_count = video_info.get('view_count', task.view_count)
                                task.upload_date = video_info.get('upload_date', task.upload_date)
                                
                                # Check for private or unavailable videos
                                if video_info.get('is_private'):
                                    raise Exception("Video is private")
                                if video_info.get('was_live') and not video_info.get('is_live'):
                                    st.warning("This was a live stream - transcription quality may vary")
                    except Exception as e:
                        st.warning(f"Failed to get video info: {str(e)}")
                        if "Video unavailable" in str(e):
                            raise  # Raise immediately for unavailable videos
                        # For other errors, continue with existing task info
                    
                    # Download audio
                    st.info(f"Downloading: {task.title} (Attempt {attempt + 1}/{max_retries})")
                    audio_info = self.manager.download_audio(
                        task.url,
                        max_retries=3,
                        retry_delay=5
                    )
                    
                    if not audio_info:
                        raise Exception("Failed to download audio")
                    
                    # Verify audio file
                    if not audio_info.get('path') or not Path(audio_info['path']).exists():
                        raise Exception("Audio file not found after download")
                    
                    # Transcription phase
                    st.info(f"Transcribing: {task.title}")
                    
                    # Check/load model
                    if not hasattr(self, 'model'):
                        st.warning("Model not loaded, loading now...")
                        if torch.cuda.is_available():
                            torch.cuda.empty_cache()
                        self.model = whisper.load_model(model_name)
                    
                    # Transcribe with error handling
                    try:
                        result = self.model.transcribe(
                            audio_info['path'],
                            language="ur",
                            task="transcribe",
                            fp16=False
                        )
                    except torch.cuda.OutOfMemoryError:
                        if torch.cuda.is_available():
                            torch.cuda.empty_cache()
                        time.sleep(retry_delay)
                        raise Exception("GPU out of memory - retrying")
                    except Exception as e:
                        if "out of memory" in str(e):
                            if torch.cuda.is_available():
                                torch.cuda.empty_cache()
                            time.sleep(retry_delay)
                            raise Exception("Memory error - retrying")
                        raise

                    # Generate transcript
                    doc_path = self.manager.save_transcript(
                        result,
                        {
                            'duration': task.duration,
                            'channel': task.channel,
                            'title': task.title,
                            'model': {
                                'name': model_name,
                                'display_name': next((k for k, v in ModelConfig.MODELS.items() 
                                                    if v['key'] == model_name), model_name),
                                'parameters': ModelConfig.MODELS.get(next((k for k, v in ModelConfig.MODELS.items() 
                                                                        if v['key'] == model_name), ''), {}).get('size', '')
                            }
                        },
                        task.url
                    )

                    if doc_path:
                        self.update_task_status(video_id, 'completed', transcript_path=doc_path)
                        task.processed_duration = task.duration
                        return {'success': True, 'path': doc_path}
                    else:
                        raise Exception("Failed to generate transcript")
                        
                except Exception as e:
                    error_msg = str(e)
                    
                    # Handle specific errors
                    if any(msg in error_msg.lower() for msg in ["video unavailable", "private video", "sign in"]):
                        self.update_task_status(video_id, 'failed', error_msg)
                        return {'success': False, 'error': error_msg}
                    
                    # Handle memory errors
                    if any(msg in error_msg.lower() for msg in ["cannot reshape tensor", "out of memory", "gpu out of memory"]):
                        if torch.cuda.is_available():
                            torch.cuda.empty_cache()
                        time.sleep(retry_delay * 2)
                        continue
                    
                    # Handle rate limiting
                    if "429" in error_msg or "too many requests" in error_msg.lower():
                        wait_time = retry_delay * (2 ** attempt)  # Exponential backoff
                        st.warning(f"Rate limit hit. Waiting {wait_time} seconds before retry...")
                        time.sleep(wait_time)
                        continue
                    
                    if attempt < max_retries - 1:
                        st.warning(f"Attempt {attempt + 1} failed for {task.title}: {error_msg}")
                        time.sleep(retry_delay)
                        continue
                    else:
                        self.update_task_status(video_id, 'failed', error_msg)
                        return {'success': False, 'error': error_msg}
                        
                finally:
                    # Cleanup audio file
                    if audio_info and 'path' in audio_info:
                        try:
                            Path(audio_info['path']).unlink(missing_ok=True)
                        except Exception as e:
                            st.warning(f"Failed to cleanup audio file: {str(e)}")

            except Exception as e:
                if attempt < max_retries - 1:
                    st.warning(f"Attempt {attempt + 1} failed: {str(e)}")
                    time.sleep(retry_delay)
                    continue
                else:
                    error_msg = str(e)
                    self.update_task_status(video_id, 'failed', error_msg)
                    return {'success': False, 'error': error_msg}

        final_error = 'Max retries exceeded'
        self.update_task_status(video_id, 'failed', final_error)
        return {'success': False, 'error': final_error}

    def display_status_table(self, container):
        """Display status table with action buttons"""
        with container:
            st.markdown("### 📊 Processing Status")
            
            # Create header for the table
            header_cols = st.columns([0.5, 2, 1, 1, 1, 1.5, 1])
            with header_cols[0]:
                st.markdown("**#**")
            with header_cols[1]:
                st.markdown("**Title**")
            with header_cols[2]:
                st.markdown("**Duration**")
            with header_cols[3]:
                st.markdown("**Status**")
            with header_cols[4]:
                st.markdown("**Progress**")
            with header_cols[5]:
                st.markdown("**Last Update**")
            with header_cols[6]:
                st.markdown("**Actions**")
            
            st.markdown("---")
            
            timestamp = datetime.now().strftime("%H%M%S")
            failed_tasks = []
            
            # Display each video's status
            for video_id, task in sorted(self.tasks.items(), key=lambda x: x[1].idx):
                cols = st.columns([0.5, 2, 1, 1, 1, 1.5, 1])
                
                # Index
                with cols[0]:
                    st.text(f"{task.idx}")
                
                # Title with truncation
                with cols[1]:
                    full_title = task.title
                    display_title = full_title[:40] + "..." if len(full_title) > 40 else full_title
                    st.markdown(f"{display_title}")
                
                # Duration
                with cols[2]:
                    st.text(format_duration(task.duration))
                
                # Status with color coding
                with cols[3]:
                    status_colors = {
                        'pending': ('🔄', 'gray'),
                        'downloading': ('⬇️', 'blue'),
                        'transcribing': ('🔄', 'blue'),
                        'completed': ('✅', 'green'),
                        'failed': ('❌', 'red')
                    }
                    emoji, color = status_colors.get(task.status, ('❓', 'gray'))
                    st.markdown(f"<span style='color: {color}'>{emoji} {task.status}</span>", 
                            unsafe_allow_html=True)
                
                # Progress
                with cols[4]:
                    if task.status != 'pending':
                        st.progress(task.processing_progress)
                    else:
                        st.text("Waiting...")
                
                # Last update time
                with cols[5]:
                    if task.last_attempt:
                        try:
                            last_time = datetime.fromisoformat(task.last_attempt)
                            time_str = last_time.strftime("%H:%M:%S")
                            st.text(time_str)
                        except:
                            st.text("--:--:--")
                    else:
                        st.text("--:--:--")
                
                # Action buttons
                with cols[6]:
                    if task.status == 'completed' and task.transcript_path:
                        try:
                            with open(task.transcript_path, 'rb') as f:
                                st.download_button(
                                    "📄",
                                    f,
                                    file_name=Path(task.transcript_path).name,
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key=f"download_{video_id}_{timestamp}"
                                )
                        except Exception as e:
                            st.error("⚠️ Error loading file")
                    
                    elif task.status == 'failed':
                        failed_tasks.append(video_id)
                        if st.button("🔄",
                                key=f"retry_{video_id}_{timestamp}"):
                            st.info(f"Retrying: {task.title}")
                            self.process_playlist(None, "base", [video_id])
            
            # Summary statistics at the bottom
            st.markdown("---")
            stats = self.get_task_statistics()
            
            cols = st.columns(4)
            with cols[0]:
                st.metric("Completed", f"{stats['completed']}/{stats['total']}")
            with cols[1]:
                st.metric("Failed", str(stats['failed']))
            with cols[2]:
                st.metric("Processing", str(stats['processing']))
            with cols[3]:
                st.metric("Pending", str(stats['pending']))
            
            # Controls for failed tasks
            if failed_tasks:
                st.markdown("---")
                control_cols = st.columns([2, 1])
                with control_cols[0]:
                    st.error(f"🚨 {len(failed_tasks)} videos failed to process")
                with control_cols[1]:
                    if st.button("🔄 Retry All Failed",
                            key=f"retry_all_failed_{timestamp}"):
                        st.info("Retrying all failed videos...")
                        self.process_playlist(None, "base", failed_tasks)

    def move_failed_audio(self, video_id, audio_path):
        """Move failed audio file to failed directory"""
        if audio_path and Path(audio_path).exists():
            try:
                failed_dir = self.current_playlist_dir / "failed"
                failed_path = failed_dir / f"{video_id}.mp3"
                shutil.move(audio_path, str(failed_path))
                return str(failed_path)
            except Exception as e:
                st.warning(f"Failed to move failed audio file: {str(e)}")
        return None

    def update_task_status(self, video_id: str, status: str, error: Optional[str] = None, 
                          transcript_path: Optional[str] = None, audio_path: Optional[str] = None):
        """Update task status with thread safety"""
        with self.lock:
            if video_id in self.tasks:
                task = self.tasks[video_id]
                task.status = status
                if error:
                    task.error = error
                if transcript_path:
                    task.transcript_path = transcript_path
                if audio_path:
                    task.audio_path = audio_path
                task.last_attempt = datetime.now().isoformat()


    def create_summary(self, playlist_info: Dict[str, Any], output_dir: Path, model_name: str):
        """Create comprehensive summary with error details and better error handling"""
        try:
            doc = Document()
            
            # Title - with safe handling
            title = doc.add_heading(f"Playlist: {playlist_info.get('title', 'Unknown Playlist')}", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Summary info
            doc.add_heading('Playlist Information', level=1)
            summary_table = doc.add_table(rows=8, cols=2)  # Added row for model info
            summary_table.style = 'Table Grid'
            
            # Get statistics and error summary
            try:
                stats = self.get_task_statistics()
            except Exception as e:
                st.warning(f"Failed to get statistics: {str(e)}")
                stats = {
                    'total': len(self.tasks),
                    'completed': 0,
                    'failed': 0,
                    'total_duration': 0,
                    'processed_duration': 0
                }

            try:
                error_summary = self.get_error_summary()
            except Exception as e:
                st.warning(f"Failed to get error summary: {str(e)}")
                error_summary = {'counts': {}, 'details': []}
            
            # Get model display info with safe handling
            try:
                model_display_name = None
                model_params = None
                for display_name, config in ModelConfig.MODELS.items():
                    if config['key'] == model_name:
                        model_display_name = display_name
                        model_params = config['size']
                        break
                
                if not model_display_name:
                    model_display_name = model_name.title()
                    model_params = "Unknown parameters"
            except Exception:
                model_display_name = model_name
                model_params = "Unknown parameters"
            
            # Fill summary table with safe handling
            try:
                rows = summary_table.rows
                rows[0].cells[0].text = 'Model Used'
                rows[0].cells[1].text = f"{model_display_name} ({model_params})"
                
                rows[1].cells[0].text = 'Total Videos'
                rows[1].cells[1].text = str(stats.get('total', 0))
                
                rows[2].cells[0].text = 'Successfully Processed'
                rows[2].cells[1].text = str(stats.get('completed', 0))
                
                rows[3].cells[0].text = 'Failed Videos'
                rows[3].cells[1].text = str(stats.get('failed', 0))
                
                rows[4].cells[0].text = 'Total Duration'
                rows[4].cells[1].text = format_duration(stats.get('total_duration', 0))
                
                rows[5].cells[0].text = 'Processed Duration'
                rows[5].cells[1].text = format_duration(stats.get('processed_duration', 0))
                
                rows[6].cells[0].text = 'Success Rate'
                total = stats.get('total', 0)
                completed = stats.get('completed', 0)
                success_rate = (completed / total * 100) if total > 0 else 0
                rows[6].cells[1].text = f"{success_rate:.1f}%"
                
                rows[7].cells[0].text = 'Processing Date'
                rows[7].cells[1].text = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            except Exception as e:
                st.warning(f"Error filling summary table: {str(e)}")
            
            # Add error summary with safe handling
            if error_summary.get('counts'):
                try:
                    doc.add_heading('Error Summary', level=1)
                    error_table = doc.add_table(rows=len(error_summary['counts']), cols=2)
                    error_table.style = 'Table Grid'
                    
                    for idx, (error_type, count) in enumerate(error_summary['counts'].items()):
                        row = error_table.rows[idx]
                        row.cells[0].text = str(error_type)
                        row.cells[1].text = str(count)
                except Exception as e:
                    st.warning(f"Error adding error summary: {str(e)}")
            
            # Add detailed error log with safe handling
            if error_summary.get('details'):
                try:
                    doc.add_heading('Error Details', level=1)
                    for error in error_summary['details']:
                        p = doc.add_paragraph()
                        p.add_run(f"Video: {error.get('title', 'Unknown')}\n").bold = True
                        p.add_run(f"Error: {error.get('error', 'Unknown error')}\n")
                        p.add_run(f"Attempts: {error.get('attempts', 0)}")
                        doc.add_paragraph()
                except Exception as e:
                    st.warning(f"Error adding error details: {str(e)}")
            
            # Save summary with safe handling
            summary_path = output_dir / "00_playlist_summary.docx"
            try:
                doc.save(str(summary_path))
                return summary_path
            except Exception as e:
                st.error(f"Failed to save summary document: {str(e)}")
                return None
                
        except Exception as e:
            st.error(f"Error creating summary: {str(e)}")
            return None
    
    def get_task_statistics(self):
        """Get comprehensive processing statistics with safe handling"""
        try:
            with self.lock:
                return {
                    'total': len(self.tasks),
                    'completed': len([t for t in self.tasks.values() if t.status == 'completed']),
                    'failed': len([t for t in self.tasks.values() if t.status == 'failed']),
                    'processing': len([t for t in self.tasks.values() 
                                    if t.status in ['processing', 'downloading', 'transcribing']]),
                    'pending': len([t for t in self.tasks.values() if t.status == 'pending']),
                    'total_duration': sum(getattr(task, 'duration', 0) for task in self.tasks.values()),
                    'processed_duration': sum(getattr(task, 'processed_duration', 0) for task in self.tasks.values())
                }
        except Exception as e:
            st.warning(f"Error getting task statistics: {str(e)}")
            return {
                'total': 0,
                'completed': 0,
                'failed': 0,
                'processing': 0,
                'pending': 0,
                'total_duration': 0,
                'processed_duration': 0
            }

    def get_error_summary(self):
        """Get comprehensive error summary"""
        error_counts = {}
        for task in self.tasks.values():
            if task.error:
                error_type = task.error.split(':')[0].strip()
                error_counts[error_type] = error_counts.get(error_type, 0) + 1
        return dict(sorted(error_counts.items(), key=lambda x: x[1], reverse=True))